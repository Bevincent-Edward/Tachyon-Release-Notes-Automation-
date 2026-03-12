from docx import Document

doc = Document(r"D:\Tachyon\Tachyon Credit,CLM,Kernel,Saturn SaaS 14th Dec'25 Release notes.docx")

print("=" * 80)
print("FULL DOCUMENT TEXT")
print("=" * 80)

full_text = []
for para in doc.paragraphs:
    full_text.append(para.text)
    
print("\n".join(full_text))

print("\n" + "=" * 80)
print("TABLE ANALYSIS")
print("=" * 80)

for i, table in enumerate(doc.tables):
    print(f"\n--- Table {i+1} ---")
    for j, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"Row {j+1}: {cells}")
